import tkinter
import Evelyn
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import easygui
# Switching to 1 enable a GUI input
guiInput = 1
startTimes = []
endTimes = []
eventTitles = []


def addHead():
    # add date heading
    if (guiInput == 0):
        date = jsonInfo['planDate'] + " PLANS"
    else:
        easygui.msgbox(
            "Welcome to SubShell - A guided process that generates a template for you to edit, and your substitute teacher to follow. Crafted with care by Adam Bloom, Chris Bonilla, Hannah Cline, and Jonathan Davis.","Welcome to SubShell")
        msg = "What day will the substitute use the plans?"
        title = "Date"
        fieldNames = ["Day of the Week", "Month (numeral)", "Day (numeral)", "Year (numeral)"]
        fieldValues = []  # we start with blanks for the values
        fieldValues = easygui.multenterbox(msg, title, fieldNames)

    # make sure that none of the fields was left blank
        while 1:
            if fieldValues == None: break
            errmsg = ""
            for i in range(len(fieldNames)):
                if fieldValues[i].strip() == "":
                    errmsg = errmsg + ('"%s" is a required field.\n\n' % fieldNames[i])
            if errmsg == "": break  # no problems found
            fieldValues = easygui.multenterbox(errmsg, title, fieldNames, fieldValues)
            print "Reply was:", fieldValues
        date = fieldValues[0] + ", " + fieldValues[1] + "/" + fieldValues[2] + "/" + fieldValues[3]

    heading = document.add_heading(date, 0)
    heading = WD_ALIGN_PARAGRAPH.CENTER
    if (guiInput == 0):
        #sub specific info , specials + customHeaderMessage
        if (jsonInfo['specialExists'] == "true"):
            subSpecific = "Our students have specials today from " + jsonInfo["specialStartTime"] + " to " + jsonInfo[
            "specialEndTime"] + ". " + jsonInfo["customHeaderMessage"]
        else:
            subSpecific = jsonInfo["customHeaderMessage"]
    else:
        msg = "What info would you like to appear at the top of the front page, below the title? For example, During recess and lunch, please check with the main office to determine if you are needed to cover indoor or outdoor recess or to support in the cafeteria."
        subSpecific = easygui.enterbox(msg,"Header Informaiton")

    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = subSpecific
    table.style = 'TableGrid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #Add whitespace
    document.add_paragraph(" ")


def trustedStudents():
    table = document.add_table(rows=2, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Classroom Helpers"
    if(guiInput==0):
        trust = jsonInfo["classroomHelpers"]
    else:
        trust = easygui.enterbox("Please enter in any information about Classroom Leaders","Classroom Leaders")
    cell = table.cell(1, 0)
    cell.text = trust
    table.style = 'TableGrid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Add whitespace
    document.add_paragraph(" ")

def teamMembers():
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Team Members"
    table.style = 'TableGrid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hasAnother = True
    count = 1
    teachers = []

    while(hasAnother is True):

        tempTeacher = easygui.enterbox('Please enter in any information about team member ' + str(count) + '.','Team Member ' + str(count))
        teachers.append(tempTeacher)
        count+=1
        hasAnother = easygui.ynbox('Would you like to add another team member?', 'Add Another Team Member', ('Yes', 'No'))

    table = document.add_table(rows=count-1, cols=1)
    x=0

    #Fill out variable length table
    while (x < count-1):
        cell = table.cell(x, 0)
        cell.text = teachers[x]
        x += 1

    table.style = 'TableGrid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    document.add_paragraph(" ")


def teacherSchedule():
    # One row for each event in the schedule plus a title row
    count = 1
    if (guiInput == 1):
        addEvent = True
        while(addEvent is True):
            msg = "Please enter in information about event " + str(count) + " in the teacher's schedule"
            title = "Event " + str(count)
            fieldNames = ["Start Time", "End Time", "Event Name"]
            fieldValues = []  # we start with blanks for the values
            fieldValues = easygui.multenterbox(msg, title, fieldNames)

    # make sure that none of the fields was left blank
            while 1:
                if fieldValues == None: break
                errmsg = ""
                for i in range(len(fieldNames)):
                    if fieldValues[i].strip() == "":
                        errmsg = errmsg + ('"%s" is a required field.\n\n' % fieldNames[i])
                if errmsg == "": break  # no problems found
                fieldValues = easygui.multenterbox(errmsg, title, fieldNames, fieldValues)
                print "Reply was:", fieldValues

            startTimes.append(fieldValues[0])
            endTimes.append(fieldValues[1])
            eventTitles.append(fieldValues[2])
            count+=1
            addEvent = easygui.ynbox('Would you like to add another event after ' + eventTitles[count-2] + '?', 'Add Another Event', ('Yes', 'No'))

        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "Teacher's Schedule"
        table.style = 'TableGrid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        #Left column time , right column event
        table = document.add_table(rows=count-1, cols=2)

        x = 0;

        #Fill out variable length table
        while (x < count-1):
            leftcell = table.cell(x, 0)
            rightcell = table.cell(x, 1)
            leftcell.text = startTimes[x] + " - " + endTimes[x]
            rightcell.text = eventTitles[x]
            x += 1

    table.style = 'TableGrid'

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #Add whitespace
    document.add_paragraph(" ")


def computerLogin():
    if(guiInput == 1):
        tusername =  easygui.enterbox("What is your computer username?","Computer Username")
        tpassword =  easygui.passwordbox("What is your computer password?","Computer Password")

    # Table header
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Computer Log-In"
    table.style = 'TableGrid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #Append 2x2 Username and Password
    table = document.add_table(rows=2, cols=2)
    cell = table.cell(0, 0)
    cell.text = "Username"
    cell = table.cell(0, 1)
    cell.text = tusername
    cell = table.cell(1, 0)
    cell.text = "Password"
    cell = table.cell(1, 1)
    cell.text = tpassword
    table.style = 'TableGrid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #Add whitespace
    document.add_paragraph(" ")


def drillInfo():
    x = 0
    document.add_heading("Fire Drill, Shelter and Lockdown information:", 1)

    # Determine how many rows in table
    if(guiInput==0):
        if (jsonInfo['fireDrillProcedures'] != ""):
            x += 1
        if (jsonInfo['lockdownProcedures'] != ""):
            x += 1
        if (jsonInfo['shelterProcedures'] != ""):
            x += 1

        table = document.add_table(rows=x, cols=2)
        count = 0

        while (count < x):
            if (jsonInfo['fireDrillProcedures'] != ""):
                cell = table.cell(count, 0)
                cell.text = "Fire Drill Procedures"
                cell = table.cell(count, 1)
                cell.text = jsonInfo['fireDrillProcedures']
                count += 1
            if (jsonInfo['lockdownProcedures'] != ""):
                cell = table.cell(count, 0)
                cell.text = "Lockdown Procedures"
                cell = table.cell(count, 1)
                cell.text = jsonInfo['lockdownProcedures']
                count += 1
            if (jsonInfo['shelterProcedures'] != ""):
                cell = table.cell(count, 0)
                cell.text = "Shelter Procedures"
                cell = table.cell(count, 1)
                cell.text = jsonInfo['shelterProcedures']
                count += 1
    else:
        table = document.add_table(rows=3, cols=2)
        cell = table.cell(0, 0)
        cell.text = "Fire Drill Procedures"
        cell = table.cell(0, 1)
        cell.text = easygui.enterbox("Please enter in any information about firedrill procedures", "Fire Drill")
        cell = table.cell(1, 0)
        cell.text = "Lockdown Procedures"
        cell = table.cell(1, 1)
        cell.text = easygui.enterbox("Please enter in any information about lockdown procedures", "Lockdown Procedures")
        cell = table.cell(2, 0)
        cell.text = "Shelter in Place Procedures"
        cell = table.cell(2, 1)
        cell.text = easygui.enterbox("Please enter in any information about shelter in place procedures", "Shelter in Place")


    table.style = 'TableGrid'

    #Add whitespace
    document.add_paragraph(" ")


def classMGMT():
    document.add_heading("Classroom Management:", 1)
    if(guiInput==0):
        mgmtText = jsonInfo['classroomManagement']
    else:
        mgmtText = easygui.enterbox("Please enter in any information about classroom management","Classroom Management")

    document.add_paragraph(mgmtText)


    # Add whitespace
    document.add_paragraph(" ")


def addPlans():
    document.add_heading("Plans:", 1)

    eventTotal = 5
    startTime = "1PM"
    endTime = "2PM"
    event = "Rotation 1"
    title = " "
    routines = " "
    x = 0

    # Repeat for each event in eventTotal
    while (x < eventTitles.__len__()):
        #Table header
        startTime = startTimes[x]
        endTime = endTimes[x]
        event = eventTitles[x]
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = " " + startTime + " to " + endTime + ": " + event
        table.style = 'TableGrid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        #Append a title with title/routines
        table = document.add_table(rows=1, cols=2)
        cell = table.cell(0, 0)
        cell.text = title
        cell = table.cell(0, 1)
        cell.text = routines
        table.style = 'TableGrid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        #Add whitespace
        document.add_paragraph(" ")
        x += 1

    #Add whitespace
    document.add_paragraph(" ")


def saveDoc(title):
    document.save(title)


def fetchJSON():
    with open('jsonOut.txt', 'r') as myfile:
        data = myfile.read()

    return str(data)


if __name__ == "__main__":
    addHead()
    teamMembers()
    trustedStudents()
    teacherSchedule()
    computerLogin()
    drillInfo()
    classMGMT()
    addPlans()
    if(guiInput==0):
        saveDoc("123testdoc.docx")
    else:
        docTitle = easygui.enterbox("Thank you for using SubShell. What would you like to name your document? Do not use any spaces or special characters.","Document Title")
        saveDoc(docTitle+".docx")